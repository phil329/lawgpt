import json
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt,RGBColor



def json2docx(json_data, doc_file):

    document = Document()

    document.styles['Normal'].font.name = u'仿宋'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'仿宋')
    document.styles['Normal'].font.size = Pt(14)
    document.styles['Normal'].font.color.rgb = RGBColor(0,0,0)


    # 标题
    # 标题等级如1,2,3这些数字，一级标题二级标题这样
    Head = document.add_heading("",level=1)# 这里不填标题内容 
    run  = Head.add_run("民事起诉状")
    Head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run.font.name=u'宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    run.font.size = Pt(22) 
    run.font.color.rgb = RGBColor(0,0,0)
    paragraph = document.add_paragraph()


    # 添加原告信息
    for plaintiff in json_data['原告']:

        ### 公司
        if '公司名称' in plaintiff:
            
            doc = document.add_paragraph()
            doc.paragraph_format.first_line_indent = Pt(14) * 2
            doc.add_run('原告：').bold = True
            doc.add_run(f"公司名称：{plaintiff['公司名称'] or '__________'}，"
                        f"地址：{plaintiff['公司所在地'] or '________________'}。"
                        f"统一社会信用代码：{plaintiff['统一社会信用代码'] or '________________'}")
            if '无' in plaintiff['法人']:
                doc.add_run('无法人')
                doc.paragraph_format.line_spacing = Pt(30) # 固定值30磅
            else:
                doc.add_run(f"法人：{plaintiff['法人']['姓名'] or '_______'}，职务：{plaintiff['法人']['职务'] or '______'}，联系方式：{plaintiff['法人']['联系方式'] or '___________'}。")
                doc.paragraph_format.line_spacing = Pt(30) # 固定值30磅
            if '无' in plaintiff['委托诉讼代理人']:
                doc = document.add_paragraph(f"无委托诉讼代理人。")
                doc.paragraph_format.first_line_indent = Pt(14) * 2
                doc.paragraph_format.line_spacing = Pt(30) # 固定值30磅
            else :
                doc = document.add_paragraph(f"委托诉讼代理人：{plaintiff['委托诉讼代理人']['姓名']}，{plaintiff['委托诉讼代理人']['事务所']}。")
                doc.paragraph_format.first_line_indent = Pt(14) * 2
                doc.paragraph_format.line_spacing = Pt(30) # 固定值30磅
        ### 个人
        else:
            doc = document.add_paragraph()
            doc.paragraph_format.first_line_indent = Pt(14) * 2
            doc.add_run('原告：').bold = True
            doc.add_run(f"{plaintiff['姓名'] or '姓名：_______'}，{plaintiff['性别'] or '性别：__'}，{plaintiff['出生日期'] or '__________'}生，"
                        f"{plaintiff['民族'] or '____' + '族'}，现住{plaintiff['住址']or '_________________'}。"
                        f"联系方式：{plaintiff['联系方式'] or '_________'}"
                        f"身份证号：{plaintiff['身份证号'] or '________________'}"
                        f"法定代理人：{plaintiff['法定代理人'] or '__________'}")
            doc.paragraph_format.line_spacing = Pt(30) # 固定值18磅
            if '无' in plaintiff['委托诉讼代理人']:
                doc = document.add_paragraph(f"无委托诉讼代理人。")
                doc.paragraph_format.first_line_indent = Pt(14) * 2
                doc.paragraph_format.line_spacing = Pt(30) # 固定值30磅
            else :
                doc = document.add_paragraph(f"委托诉讼代理人：{plaintiff['委托诉讼代理人']['姓名']}，{plaintiff['委托诉讼代理人']['事务所']}。")
                doc.paragraph_format.first_line_indent = Pt(14) * 2
                doc.paragraph_format.line_spacing = Pt(30) # 固定值30磅
    # document.add_paragraph()


    # 添加被告信息
    for defendant in json_data['被告']:

        ### 公司
        if '公司名称' in defendant:
            doc = document.add_paragraph()
            doc.paragraph_format.first_line_indent = Pt(14) * 2
            doc.add_run('被告：').bold = True
            doc.add_run(f"公司名称：{defendant['公司名称'] or '__________'}，"
                        f"地址：{defendant['公司所在地'] or '________________'}。"
                        f"统一社会信用代码：{defendant['统一社会信用代码'] or '________________'}")
            if '无' in defendant['法人']:
                doc.add_run('无法人。')
                doc.paragraph_format.line_spacing = Pt(30) # 固定值30磅
            else:
                doc.add_run(f"法人：{defendant['法人']['姓名'] or '_______'}，职务：{defendant['法人']['职务'] or '______'}，联系方式：{defendant['法人']['联系方式'] or '___________'}。")
                doc.paragraph_format.line_spacing = Pt(30) # 固定值30磅
            if '无' in defendant['委托诉讼代理人']:
                doc = document.add_paragraph(f"无委托诉讼代理人。")
                doc.paragraph_format.first_line_indent = Pt(14) * 2
                doc.paragraph_format.line_spacing = Pt(30) # 固定值30磅
            else:
                doc = document.add_paragraph(f"委托诉讼代理人：{defendant['委托诉讼代理人']['姓名']}，{defendant['委托诉讼代理人']['事务所']}。")
                doc.paragraph_format.first_line_indent = Pt(14) * 2
                doc.paragraph_format.line_spacing = Pt(30) # 固定值30磅
        ## 个人
        else:
            doc = document.add_paragraph()
            doc.paragraph_format.first_line_indent = Pt(14) * 2
            doc.add_run('被告：').bold = True
            doc.add_run(f"{defendant['姓名'] or '姓名：_______'}，{defendant['性别'] or '性别：__'}，{defendant['出生日期'] or '__________'}生，"
                        f"{defendant['民族'] or '____' + '族'}，现住{defendant['住址']or '_________________'}。"
                        f"联系方式：{defendant['联系方式'] or '_________'}"
                        f"身份证号：{defendant['身份证号'] or '________________'}"
                        f"法定代理人：{defendant['法定代理人'] or '__________'}")
            doc.paragraph_format.line_spacing = Pt(30) # 固定值18磅
            if '无' in defendant['委托诉讼代理人']:
                doc = document.add_paragraph(f"无委托诉讼代理人。")
                doc.paragraph_format.first_line_indent = Pt(14) * 2
                doc.paragraph_format.line_spacing = Pt(30) # 固定值30磅
            else:
                doc = document.add_paragraph(f"委托诉讼代理人：{defendant['委托诉讼代理人']['姓名']}，{defendant['委托诉讼代理人']['事务所']}。")
                doc.paragraph_format.first_line_indent = Pt(14) * 2
                doc.paragraph_format.line_spacing = Pt(30) # 固定值30磅

    document.add_paragraph()


    # 添加诉讼请求
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Pt(14) * 2
    paragraph.add_run('诉讼请求：').bold = True
    paragraph.paragraph_format.line_spacing = Pt(30) # 固定值30磅
    if json_data['诉讼请求'] == "null":
        doc = document.add_paragraph('无诉讼请求！\n')
        doc.paragraph_format.first_line_indent = Pt(14) * 2
        doc.paragraph_format.line_spacing = Pt(30) # 固定值30磅
    else:
        request_text = json_data['诉讼请求']
        # 根据"\n"拆分为多行
        lines = request_text.split('\n')
        # 对每一行进行处理
        for line in lines:
            paragraph = document.add_paragraph(line)
            paragraph.paragraph_format.first_line_indent = Pt(14) * 2
            paragraph.paragraph_format.line_spacing = Pt(30) # 固定值30磅
    document.add_paragraph()

    # 添加事实和理由
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Pt(14) * 2
    paragraph.add_run('事实和理由：').bold = True
    paragraph.paragraph_format.line_spacing = Pt(30) # 固定值30磅
    if json_data['事实理由'] == "null":
        doc = document.add_paragraph('无事实和理由！\n')
        doc.paragraph_format.first_line_indent = Pt(14) * 2
        doc.paragraph_format.line_spacing = Pt(30) # 固定值30磅
    else:
        request_text = json_data['事实理由']
        # 根据"\n"拆分为多行
        lines = request_text.split('\n')
        # 对每一行进行处理
        for line in lines:
            paragraph = document.add_paragraph(line)
            paragraph.paragraph_format.first_line_indent = Pt(14) * 2
            paragraph.paragraph_format.line_spacing = Pt(30) # 固定值30磅
    document.add_paragraph()

    # 添加证据和证据来源，证人姓名和住所
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Pt(14) * 2
    paragraph.add_run('证据和证据来源，证人姓名和住所：').bold = True
    paragraph.paragraph_format.line_spacing = Pt(30) # 固定值30磅
    if json_data['证据'] == "null":
        doc = document.add_paragraph('无证据和证据来源，证人姓名和住所！\n')
        doc.paragraph_format.first_line_indent = Pt(14) * 2
        doc.paragraph_format.line_spacing = Pt(30) # 固定值30磅
    else:
        request_text = json_data['证据']
        # 根据"\n"拆分为多行
        lines = request_text.split('\n')
        # 对每一行进行处理
        for line in lines:
            paragraph = document.add_paragraph(line)
            paragraph.paragraph_format.first_line_indent = Pt(14) * 2
            paragraph.paragraph_format.line_spacing = Pt(30) # 固定值30磅

    # 结尾部分
    paragraph = document.add_paragraph()
    paragraph.add_run('此致').bold = True
    paragraph.paragraph_format.first_line_indent = Pt(14) * 2
    paragraph.paragraph_format.line_spacing = Pt(30) # 固定值30磅

    paragraph = document.add_paragraph()
    paragraph.add_run(json_data['法院']).bold = True
    paragraph.paragraph_format.line_spacing = Pt(30) # 固定值30磅
    paragraph = document.add_paragraph('\n附：本起诉状副本2份')
    paragraph.paragraph_format.line_spacing = Pt(30) # 固定值30磅
                        
    # 起诉人(签名)
    paragraph = document.add_paragraph()
    paragraph.add_run('起诉人(签名)\n').bold = True
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.line_spacing = Pt(30) # 固定值30磅

    # 日期
    if json_data['日期'] is None:
        date_paragraph = document.add_paragraph()
        date_paragraph.add_run('日期：____________').bold = True
        date_paragraph.paragraph_format.line_spacing = Pt(30) # 固定值30磅
    else:
        date_paragraph = document.add_paragraph()
        date_paragraph.add_run(json_data['日期']).bold = True
        date_paragraph.paragraph_format.line_spacing = Pt(30) # 固定值30磅
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


    from utils.tools import law_claim
    claim_paragraph = document.add_paragraph(law_claim)
    claim_paragraph.paragraph_format.first_line_indent = Pt(14) * 2
    claim_paragraph.paragraph_format.line_spacing = Pt(30) # 固定值30磅

    # 保存文档
    document.save(doc_file)


if __name__ == '__main__':
    json_1 = 'person.json'

    with open(json_1, 'r',encoding='utf8') as f:
        json_data = json.loads(f.read())

    json2docx(json_data, 'demo.docx')