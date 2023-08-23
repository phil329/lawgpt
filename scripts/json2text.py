import json

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from utils.iflytek import SparkApi
from utils.prompt_config import *


def generate_docx(data,doc_file):
    # 以下密钥信息从控制台获取
    appid = "f54bca1a"  # 填写控制台中获取的 APPID 信息
    api_secret = "ZGJmMjY3YWFmNGY5N2Q1YWMwOGZlMjFh"  # 填写控制台中获取的 APISecret 信息
    api_key = "2e7516edf376713eb5d99fc812a87ed5"  # 填写控制台中获取的 APIKey 信息

    text = []

    def getText(role, content):
        jsoncon = {}
        jsoncon["role"] = role
        jsoncon["content"] = content
        text.append(jsoncon)
        return text

    def getlength(text):
        length = 0
        for content in text:
            temp = content["content"]
            leng = len(temp)
            length += leng
        return length

    def checklen(text):
        while (getlength(text) > 8000):
            del text[0]
        return text

    # 用于配置大模型版本，默认“generalv2”
    domain = "generalv2"  # v2.0版本

    # 云端环境的服务地址
    Spark_url = "ws://spark-api.xf-yun.com/v2.1/chat"  # v2.0环境的地址

    # 提取原告和被告部分
    result = {
        "原告": data["原告"],
        "被告": data["被告"]
    }

    # 转换为JSON格式字符串
    result_json = json.dumps(result, ensure_ascii=False, indent=4)
    prompt = result_json + json_to_text
    question = checklen(getText("user", prompt))
    SparkApi.answer = ""
    SparkApi.main(appid, api_key, api_secret, Spark_url, domain, question)
    answer = SparkApi.answer

    document = Document()

    document.styles['Normal'].font.name = u'宋体'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    document.styles['Normal'].font.size = Pt(10.5)
    document.styles['Normal'].font.color.rgb = RGBColor(0,0,0)


    # 标题
    document.add_heading('民事起诉状', level=0).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    document.add_paragraph(answer)

    # 添加诉讼请求
    document.add_paragraph('诉讼请求：\n' + data['诉讼请求'])

    # 添加事实和理由
    document.add_paragraph('事实和理由：\n' + data['事实理由'])

    # 添加证据和证据来源，证人姓名和住所
    document.add_paragraph('证据和证据来源，证人姓名和住所：\n' + data['证据'])

    # 结尾部分
    document.add_paragraph('此致\n\n' + data['法院'] + '人民法院\n\n附：本起诉状副本×份')

    # 起诉人(签名)
    paragraph = document.add_paragraph('起诉人(签名)\n')
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # 日期
    date_paragraph = document.add_paragraph(data['日期'])
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # 保存文档为docx文件
    document.save(doc_file)

if __name__ == "__main__":
    # 从JSON文件中读取数据
    doc_file = './legal_document02.docx'

    with open('./person.json', 'r', encoding='utf-8') as f:
        data = json.load(f)
    docx_file = generate_docx(data,doc_file)